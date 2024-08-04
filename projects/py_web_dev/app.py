import os
from urllib.parse import quote
from __init__ import app, db, general, IMAGE_DIR, PATH_STATIC, PATH_APP
import lib.db_models as models

from flask import Flask, render_template, jsonify, request

from docx import Document
from docx.shared import Inches, Pt, RGBColor


z0 = general (id0=0, id1=22, s0="user0", s1="email0")
z1 = general (id0=1, id1=44, s0="user1", s1="email1")
tag1, tag2 = 1, 10


@app.route('/')
def root():
    return 'This is root page'

@app.route('/prx')
def prxe():
    if (db):
        db.drop_all()
    db.create_all()

    u0 = models.User (id=z0.id0, username=z0.s0, email=z0.s1)
    u1 = models.User (id=z1.id0, username=z1.s0, email=z1.s1)
    db.session.add(u0)
    db.session.add(u1)
    db.session.commit()

    all_rows = "\n<br>".join (f"<b>[{x.id}, {x.email}, {x.username}]</b>" for x in models.User.query.all())
    #print (all_rows)
    return (f"<h1>This is prx page.</h1> <p>{all_rows}</p>")


@app.route('/crx')
def crx():
    return render_template ("crx2.html")


@app.route ("/xrv", methods=["POST"])
def post ():
    data = request.get_json ()

    if (db):
        db.drop_all()
    db.create_all()

    u0 = models.User(id=data.get("id0"), username=data.get("username0"), email=data.get("email0"))
    u1 = models.User(id=data.get("id1"), username=data.get("username1"), email=data.get("email1"))
    db.session.add(u0)
    db.session.add(u1)
    db.session.commit()

    all_rows = "\n<br>".join(f"<b>[{x.id}, {x.email}, {x.username}]</b>" for x in models.User.query.all())
    # print (all_rows)
    return (f"<h1>This is prx page.</h1> <p>{all_rows}</p>")



@app.route ("/comm", methods=["GET", "POST"])
def compile_image ():
    image_list_rel = [os.path.join ("/static/image/", im) for im in os.listdir(IMAGE_DIR) if im.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
    image_list_rel = sorted (image_list_rel)
    image_list_abs = [os.path.join(IMAGE_DIR, im) for im in os.listdir(IMAGE_DIR) if
                      im.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]

    #print (image_list_rel)
    doc = Document()
    section = doc.sections[0]
    margins = {
        'top': 0.4,  # 1 inch
        'bottom': 0.4,  # 1 inch
        'left': 0.4,  # 1 inch
        'right': 0.4  # 1 inch
    }
    section.top_margin = Inches(margins['top'])
    section.bottom_margin = Inches(margins['bottom'])
    section.left_margin = Inches(margins['left'])
    section.right_margin = Inches(margins['right'])


    if (request.method=="GET"):
        return render_template("index_im.html", images=image_list_rel)

    if (request.method=="POST"):
        global tag1, tag2

        for im in image_list_rel:
            #doc.add_heading (request.form.get (im, ""), level=1)
            #print (PATH_APP[:-1]+im)

            para = doc.add_heading(level=1)
            run = para.add_run (request.form.get (im, ""))
            run.font.size = Pt (20)
            run.font.color.rgb = RGBColor (0, 0, 0)
            run.font.bold = True

            doc.add_picture (PATH_APP[:-1]+im, width=Inches (3.0))

            if (db):
                db.drop_all()
            db.create_all()

            new_entry = models.image_title (
                tag1=tag1,
                tag2=tag2,
                image_title_value=request.form.get (im, ""),
                image_name=im,
                image_abs_path=PATH_APP[:-1]+im
            )

            tag1, tag2 = tag1+1, tag2+10
            db.session.add(new_entry)

        db.session.commit()
        all_rows = "\n<br>".join(f"<b>[{x.id}, {x.image_title_value}, {x.image_name}, {x.image_abs_path}]</b>" for x in models.image_title.query.all())
        #print (all_rows)

        doc.save (PATH_STATIC+"generated_file.docx")

        return render_template ("results_im.html")


@app.route ("/comm/dash", methods=["GET"])
def image_dashboard ():
    all_rows = models.image_title.query.all()
    object_to_dict = models.image_title_auto_schema (many=True)
    json_response = jsonify (object_to_dict.dump (all_rows))
    #print (json_response)
    return json_response



if (__name__=="__main__"):
    app.run (debug=True, port=8156, host="127.0.0.1")
