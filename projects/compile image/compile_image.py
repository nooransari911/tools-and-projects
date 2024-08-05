import os, glob
import subprocess

from flask import Flask, render_template, jsonify, request, Request

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from gdrive_im import download_files
from gdrive_te import download_filest

# ROOT="/home/ansarimn/Downloads/tools and projects/projects/compile image/"
# ROOT = "/tmp/"
ROOT="/opt/render/project/src/projects/compile image/"


inits = [
    "mkdir -p {ROOT}",
    "cd {ROOT}",
    "mkdir -p ./static",
    "mkdir -p ./static/image",
    "mkdir -p ./templates"
]

[subprocess.Popen (x, shell=True).wait() for x in inits]








PATH_APP = ROOT
PATH_STATIC = PATH_APP+"static/"
PATH_TEMPLATE = PATH_APP+"templates/"
IMAGE_DIR = PATH_STATIC+"image/"


app = Flask(__name__, static_folder=PATH_STATIC, template_folder=PATH_TEMPLATE)


def gdrive ():
    #inits = "python3 ./gdrive_im.py && python3 ./gdrive_te.py"
    #subprocess.Popen (inits, shell=True)
    download_files()
    download_filest()


@app.route ("/refresh")
def refresh ():
    # [os.remove(f) for f in glob.glob(f'{IMAGE_DIR}*') if os.path.isfile(f)]
    # [os.remove(f) for f in os.listdir(IMAGE_DIR)]
    if (os.listdir (IMAGE_DIR)):
        inits = f"rm -f {IMAGE_DIR}*"
        subprocess.Popen (inits, shell=True)

    inits = "python3 ./gdrive_im.py && python3 ./gdrive_te.py"
    subprocess.Popen (inits, shell=True)
    return ("<h1>Download was success</h1>")


@app.route ("/comm", methods=["GET", "POST"])
def compile_image ():
    # gdrive()
    # print ("return from refresh")
    image_list_rel = [os.path.join ("/static/image/", im) for im in os.listdir(IMAGE_DIR) if im.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
    image_list_rel = sorted (image_list_rel)
    print (image_list_rel)
    image_list_abs = [os.path.join(IMAGE_DIR, im) for im in os.listdir(IMAGE_DIR) if
                      im.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]

    doc = Document()
    section = doc.sections[0]
    margins = {
        'top': 0.4,  # 1 inch
        'bottom': 0.4,  # 1 inch
        'left': 0.4,  # 1 inch
        'right': 0.4  # 1 inch
    }
    section.top_margin = Inches(margins['top'])
    section.bottom_margin = Inches(margins['bottom'])
    section.left_margin = Inches(margins['left'])
    section.right_margin = Inches(margins['right'])


    if (request.method=="GET"):
        return render_template("index_im.html", images=image_list_rel)

    if (request.method=="POST"):
        for im in image_list_rel:

            para = doc.add_heading(level=1)
            run = para.add_run (request.form.get (im, ""))
            run.font.size = Pt (20)
            run.font.color.rgb = RGBColor (0, 0, 0)
            run.font.bold = True

            doc.add_picture (PATH_APP[:-1]+im, width=Inches (3.0))


        doc.save (PATH_STATIC+"generated_file.docx")

        return render_template ("results_im.html")




if (__name__=="__main__"):
    app.run (debug=False)
